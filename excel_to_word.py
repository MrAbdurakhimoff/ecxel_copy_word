import pandas as pd
import docx

# Excel file o'qildi !
data = pd.read_excel('excel.xlsx')

# Word document yaratildi
doc = docx.Document()

# Paragraph qo'shildi
p = doc.add_paragraph()

# Paragraph format berildi !
p.paragraph_format.line_spacing = 1
p.paragraph_format.space_after = 0

# Paragraphga malumot qoshildi !
run = p.add_run(str(data))

# Malumot uchun format berildi !
run.bold = True
run.italic = True
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(16)

# Document "word.docx" ushbu nom bilan saqlandi !
doc.save("word.docx")